# services/resume_parser.py
import io
import re
from typing import Optional

import pdfplumber
from docx import Document as DocxDocument
from utils.helpers import extract_email, extract_phone, clean_text


def parse_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"PDF parse error: {e}")
    return text.strip()


def parse_docx(file_bytes: bytes) -> str:
    text = ""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
    except Exception as e:
        print(f"DOCX parse error: {e}")
    return text.strip()


def extract_name_heuristic(text: str) -> str:
    """First non-empty line that looks like a name (2-4 title-case words)."""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    for line in lines[:8]:
        if re.search(r'[@\d\/\|http]', line):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w):
            return line
    return ""


def extract_experience_years_robust(text: str) -> float:
    """
    Robust multi-pattern experience extraction.
    Handles: "3 years", "3+ years", "2-4 years", "Experience: 3 years",
             "3 years of experience", "three years", "3 yrs"
    """
    # Word-to-number for written-out years
    word_nums = {
        "one": 1, "two": 2, "three": 3, "four": 4, "five": 5,
        "six": 6, "seven": 7, "eight": 8, "nine": 9, "ten": 10,
    }

    patterns = [
        # "Experience: 3 years" or "Experience: 3+ years"
        r'experience\s*[:\-]\s*(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)',
        # "3 years of experience" / "3+ years of experience"
        r'(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\s+(?:of\s+)?(?:work\s+)?experience',
        # "experience of 3 years"
        r'experience\s+of\s+(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)',
        # range "2-4 years" → take lower bound
        r'(\d+)\s*[-–]\s*\d+\s*(?:years?|yrs?)',
        # "3+ years" standalone
        r'(\d+(?:\.\d+)?)\s*\+\s*(?:years?|yrs?)',
        # plain "3 years" anywhere
        r'(\d+(?:\.\d+)?)\s+(?:years?|yrs?)',
        # "X years" where X is a word
        r'(one|two|three|four|five|six|seven|eight|nine|ten)\s+(?:years?|yrs?)',
        # total experience section header
        r'total\s+experience\s*[:\-]\s*(\d+(?:\.\d+)?)',
    ]

    for p in patterns:
        m = re.search(p, text, re.IGNORECASE)
        if m:
            val = m.group(1)
            if val.lower() in word_nums:
                return float(word_nums[val.lower()])
            try:
                return float(val)
            except ValueError:
                continue

    return 0.0


def extract_education(text: str) -> str:
    """Extract highest education qualification from resume text."""
    edu_patterns = [
        r"Ph\.?D\.?[^\n,\.]{0,60}",
        r"M\.?Tech[^\n,\.]{0,60}",
        r"M\.?E\.?\b[^\n,\.]{0,60}",
        r"Master[^\n,\.]{0,60}",
        r"M\.?S\.?\b[^\n,\.]{0,60}",
        r"MBA[^\n,\.]{0,60}",
        r"M\.?Sc[^\n,\.]{0,60}",
        r"B\.?Tech[^\n,\.]{0,60}",
        r"B\.?E\.?\b[^\n,\.]{0,60}",
        r"Bachelor[^\n,\.]{0,60}",
        r"B\.?S\.?\b[^\n,\.]{0,60}",
        r"B\.?Sc[^\n,\.]{0,60}",
    ]
    for pat in edu_patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(0).strip()
    return ""


class ParsedResume:
    def __init__(self):
        self.raw_text: str = ""
        self.name: str = ""
        self.email: str = ""
        self.phone: str = ""
        self.experience_years: float = 0.0


def parse_resume_file(file_bytes: bytes, filename: str) -> ParsedResume:
    result = ParsedResume()

    ext = filename.lower().split(".")[-1]
    if ext == "pdf":
        result.raw_text = parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        result.raw_text = parse_docx(file_bytes)
    else:
        result.raw_text = file_bytes.decode("utf-8", errors="ignore")

    result.name = extract_name_heuristic(result.raw_text)
    result.email = extract_email(result.raw_text)
    result.phone = extract_phone(result.raw_text)
    result.experience_years = extract_experience_years_robust(result.raw_text)

    return result